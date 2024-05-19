from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Create a new Document
doc = Document()
doc.add_heading('Non-Disclosure Agreement', 0)

# Insert Agreement Date and Parties
doc.add_paragraph("This Agreement is made on [Date]")
doc.add_paragraph("\nBetween\n")
doc.add_paragraph("ABC (hereinafter referred to as \"Disclosing Party\"), a company organized and existing under the laws of [Jurisdiction], with its principal office located at:\n")
doc.add_paragraph("[Full Address of ABC]").paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

doc.add_paragraph("\nAnd\n")
doc.add_paragraph("XYZ or XYZorg (hereinafter referred to as \"Receiving Party\"), [Individual/Corporation] organized and existing under the laws of [Jurisdiction], with its principal office located at:\n")
doc.add_paragraph("[Full Address of XYZ/XYZorg]").paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

doc.add_paragraph("\nWhereas the Disclosing Party agrees to disclose certain proprietary information under the terms and conditions described below.\n")

# Define Sections
sections = [
    "1. Definition of Confidential Information",
    "For the purposes of this Agreement, \"Confidential Information\" shall include all information or material that has or could have commercial value or other utility in the business in which Disclosing Party is engaged. This includes, but is not limited to, codebase, documentation, business strategies, business logic, customer lists, and data.",
    "2. Non-use and Non-disclosure",
    "The Receiving Party agrees not to use the Confidential Information for any purpose except to evaluate and engage in discussions concerning a potential business relationship with the Disclosing Party. The Receiving Party shall not disclose any Confidential Information to third parties or to employees, except as required in the course of their business relationship.",
    "3. Restrictions on Copying",
    "The Receiving Party shall not copy, replicate, or reverse-engineer any part of the Confidential Information provided by the Disclosing Party without the written permission of the Disclosing Party.",
    "4. Term",
    "The confidentiality obligations of this Agreement shall remain in effect for a period of [number] years after the date of disclosure of the Confidential Information. Upon termination or expiration of this Agreement, Receiving Party shall cease use of and return or destroy all Confidential Information received under this Agreement.",
    "5. No License",
    "Nothing contained herein shall be construed as granting or conferring any rights by license or otherwise in any Confidential Information disclosed under this Agreement.",
    "6. Governing Law",
    "This Agreement and any action related thereto shall be governed by the laws of [Jurisdiction] without regard to its conflict of laws provisions.",
    "7. Entire Agreement",
    "This Agreement constitutes the entire agreement between the parties with respect to the subject matter hereof and supersedes all prior and contemporaneous oral or written agreements."
]

# Add Sections to Document
for section in sections:
    if section.startswith("1.") or section.startswith("2.") or section.startswith("3.") or section.startswith("4.") or section.startswith("5.") or section.startswith("6.") or section.startswith("7.") :
        doc.add_heading(section, level=1)
    else:
        doc.add_paragraph(section)

# Signature Section
doc.add_page_break()
doc.add_paragraph("IN WITNESS WHEREOF, the Parties have executed this Agreement as of the date first above written.\n")
doc.add_paragraph("ABC\n\nBy: ___________________________\n\nName:\n\nTitle:\n\nDate:\n")
doc.add_paragraph("XYZ/XYZorg\n\nBy: ___________________________\n\nName:\n\nTitle:\n\nDate:\n")

# Save the document
file_path = "/home/gsaravanan/Documents/NDA_Document_ABC_XYZ.docx"
doc.save(file_path)

